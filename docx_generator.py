from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

def add_spacer(doc, pt=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)

def build_classic_docx(data):
    doc = Document()
    
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(data.get('name', 'Your Name'))
    name_run.font.size = Pt(24)
    name_run.bold = True
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_p = doc.add_paragraph()
    contact_p.add_run(data.get('contact', ''))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_spacer(doc, 10)

    sections = [
        ("Summary", data.get('summary', '')),
        ("Experience", data.get('experience', [])),
        ("Skills", ", ".join(data.get('skills', []))),
        ("Projects", data.get('projects', [])),
        ("Education", data.get('education', [])),
        ("Certifications", data.get('certifications', []))
    ]

    for title, content in sections:
        if content:
            # Heading
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(title.upper())
            h_run.bold = True
            h_run.font.size = Pt(14)
            h_p.paragraph_format.space_before = Pt(12)
            h_p.paragraph_format.space_after = Pt(4)
            
            # Simple line: a sequence of underscores or a border
            border_p = doc.add_paragraph()
            border_run = border_p.add_run("_" * 70)
            # Make the border line smaller and closer to text
            border_run.font.size = Pt(8)
            border_p.paragraph_format.space_before = Pt(0)
            border_p.paragraph_format.space_after = Pt(8)

            if isinstance(content, str):
                p = doc.add_paragraph(content)
                p.paragraph_format.space_after = Pt(10)
            elif title == "Certifications":
                for c in content:
                    p = doc.add_paragraph(f"• {c}", style='List Bullet')
                    p.paragraph_format.space_after = Pt(4)
            else:
                for item in content:
                    if title == "Experience":
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('role')} | {item.get('company')}").bold = True
                        date_run = p.add_run(f"  ({item.get('date')})")
                        date_run.font.color.rgb = RGBColor(100, 100, 100)
                        
                        desc_p = doc.add_paragraph(item.get('desc', ''))
                        desc_p.paragraph_format.space_after = Pt(8)
                        
                    elif title == "Projects":
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('title')}").bold = True
                        desc_p = doc.add_paragraph(item.get('desc', ''))
                        desc_p.paragraph_format.space_after = Pt(8)
                        
                    elif title == "Education":
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('inst')} - {item.get('degree')} ({item.get('year')})").bold = True
                        if item.get('details'):
                            desc_p = doc.add_paragraph(item.get('details'))
                            desc_p.paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def build_monochrome_docx(data):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(data.get('name', 'Your Name').upper())
    name_run.font.size = Pt(22)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(17, 17, 17)
    
    contact_p = doc.add_paragraph()
    c_run = contact_p.add_run(data.get('contact', ''))
    c_run.font.color.rgb = RGBColor(102, 102, 102)
    contact_p.paragraph_format.space_after = Pt(20)

    sections = [
        ("Summary", data.get('summary', '')),
        ("Experience", data.get('experience', [])),
        ("Skills", ", ".join(data.get('skills', []))),
        ("Projects", data.get('projects', [])),
        ("Education", data.get('education', [])),
        ("Certifications", data.get('certifications', []))
    ]

    for title, content in sections:
        if content:
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(f" {title.upper()} ")
            h_run.bold = True
            h_run.font.size = Pt(12)
            # Background shading for heading requires XML manipulation, we'll keep it simple bold
            h_p.paragraph_format.space_before = Pt(16)
            h_p.paragraph_format.space_after = Pt(6)

            if isinstance(content, str):
                p = doc.add_paragraph(content)
                p.paragraph_format.space_after = Pt(10)
            elif title == "Certifications":
                for c in content:
                    p = doc.add_paragraph(f"• {c}")
                    p.paragraph_format.space_after = Pt(4)
            else:
                for item in content:
                    if title == "Experience":
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.autofit = False
                        tbl.columns[0].width = Inches(5.0)
                        tbl.columns[1].width = Inches(1.5)
                        
                        r0 = tbl.rows[0].cells
                        
                        p1 = r0[0].paragraphs[0]
                        p1.add_run(f"{item.get('role')} | {item.get('company')}").bold = True
                        p1.paragraph_format.space_after = Pt(0)
                        
                        p2 = r0[1].paragraphs[0]
                        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        d_run = p2.add_run(item.get('date', ''))
                        d_run.font.color.rgb = RGBColor(100, 100, 100)
                        p2.paragraph_format.space_after = Pt(0)
                        
                        desc_p = doc.add_paragraph(item.get('desc', ''))
                        desc_p.paragraph_format.space_after = Pt(12)
                        
                    elif title == "Projects":
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('title')}").bold = True
                        desc_p = doc.add_paragraph(item.get('desc', ''))
                        desc_p.paragraph_format.space_after = Pt(12)
                        
                    elif title == "Education":
                        p = doc.add_paragraph()
                        p.add_run(f"{item.get('inst')} - {item.get('degree')} ({item.get('year')})").bold = True
                        if item.get('details'):
                            desc_p = doc.add_paragraph(item.get('details'))
                            desc_p.paragraph_format.space_after = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def build_sidebar_docx(data):
    # Simulating a sidebar with a 1x2 table without borders
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Set margins
    sections_doc = doc.sections
    for section in sections_doc:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.5)  # Left column
    table.columns[1].width = Inches(5.0)  # Right column
    
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # --- LEFT COLUMN ---
    # In a real Word document, setting background color for the entire cell height is tricky
    # We will simulate the look with text and colors
    l_name = left_cell.add_paragraph()
    l_name.add_run(data.get('name', 'Your Name')).bold = True
    l_name.runs[0].font.size = Pt(16)
    l_name.paragraph_format.space_after = Pt(12)
    
    l_contact_h = left_cell.add_paragraph()
    l_contact_h.add_run("CONTACT").bold = True
    l_contact_h.paragraph_format.space_before = Pt(12)
    
    parts = data.get('contact', '').split('|')
    for p in parts:
        left_cell.add_paragraph(p.strip())
        
    l_skill_h = left_cell.add_paragraph()
    l_skill_h.add_run("SKILLS").bold = True
    l_skill_h.paragraph_format.space_before = Pt(12)
    left_cell.add_paragraph(", ".join(data.get('skills', [])))
    
    if data.get('certifications'):
        l_cert_h = left_cell.add_paragraph()
        l_cert_h.add_run("CERTIFICATIONS").bold = True
        l_cert_h.paragraph_format.space_before = Pt(12)
        for c in data.get('certifications'):
            left_cell.add_paragraph(f"• {c}")

    # --- RIGHT COLUMN ---
    # Helper for right headings
    def add_right_heading(text):
        p = right_cell.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(44, 62, 80)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    if data.get('summary'):
        add_right_heading("SUMMARY")
        right_cell.add_paragraph(data.get('summary', ''))
        
    if data.get('experience'):
        add_right_heading("EXPERIENCE")
        for item in data.get('experience', []):
            ptbl = right_cell.add_table(rows=1, cols=2)
            ptbl.autofit = False
            ptbl.columns[0].width = Inches(3.5)
            ptbl.columns[1].width = Inches(1.5)
            r0 = ptbl.rows[0].cells
            
            p1 = r0[0].paragraphs[0]
            p1.add_run(f"{item.get('role')} | {item.get('company')}").bold = True
            
            p2 = r0[1].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.add_run(item.get('date', '')).font.color.rgb = RGBColor(127, 140, 141)
            
            right_cell.add_paragraph(item.get('desc', ''))
            add_spacer(right_cell, 8)
            
    if data.get('projects'):
        add_right_heading("PROJECTS")
        for item in data.get('projects', []):
            p = right_cell.add_paragraph()
            p.add_run(item.get('title', '')).bold = True
            right_cell.add_paragraph(item.get('desc', ''))
            add_spacer(right_cell, 8)
            
    if data.get('education'):
        add_right_heading("EDUCATION")
        for item in data.get('education', []):
            p = right_cell.add_paragraph()
            p.add_run(f"{item.get('inst')} - {item.get('degree')} ({item.get('year')})").bold = True
            if item.get('details'):
                 right_cell.add_paragraph(item.get('details'))
            add_spacer(right_cell, 8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
